"""
hujjat_generator.py — Word (.docx) + ZIP/RAR arxiv tizimi
Rasmli, jadvallu, bo'limlarga ajratilgan hujjatlar
"""
import os, io, zipfile, psycopg2
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATABASE_URL = os.getenv("DATABASE_URL")

def db():
    return psycopg2.connect(DATABASE_URL)

# ── Stil yordamchilar ──
def set_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Header background
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "2E86AB")
        tcPr.append(shd)
        # Oq matn
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    # Qatorlar
    for r_idx, row in enumerate(rows):
        trow = table.rows[r_idx+1]
        bg   = "EBF5FB" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = trow.cells[c_idx]
            cell.text = str(val) if val is not None else ""
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Alternating color
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  bg)
            tcPr.append(shd)
    return table

def add_image_safe(doc, image_data: bytes, ext: str = "png", width_cm: float = 14.0):
    """Rasmni xavfsiz qo'shadi."""
    try:
        img_stream = io.BytesIO(image_data)
        doc.add_picture(img_stream, width=Cm(width_cm))
    except Exception as e:
        doc.add_paragraph(f"[Rasm yuklanmadi: {e}]")

def page_break(doc):
    doc.add_page_break()

# ══════════════════════════════════════
# 1. TEST NATIJALARI WORD
# ══════════════════════════════════════
def test_results_docx(fan=None, sinf=None, days=30) -> io.BytesIO:
    conn = db(); cur = conn.cursor()
    cur.execute("""
        SELECT u.full_name, u.class,
               COALESCE(d.subject_name,'—') as fan,
               COUNT(*) as jami,
               SUM(CASE WHEN th.correct THEN 1 ELSE 0 END) as togri,
               ROUND(100.0*SUM(CASE WHEN th.correct THEN 1 ELSE 0 END)/COUNT(*)) as pct
        FROM test_history th
        JOIN users u ON u.user_id = th.user_id
        LEFT JOIN dts_tree d ON d.topic_code = th.topic_code
        WHERE th.created_at >= NOW() - INTERVAL '%s days'
        GROUP BY u.full_name, u.class, d.subject_name
        ORDER BY pct DESC
    """, (days,))
    rows = cur.fetchall()
    cur.close(); conn.close()

    doc  = Document()
    doc.core_properties.author = "SamTM Ta'lim Tizimi"

    set_heading(doc, "Test Natijalari Hisoboti", 1, (46,134,171))
    p = doc.add_paragraph()
    p.add_run(f"Sana: {datetime.now().strftime('%d.%m.%Y')}").italic = True
    if fan:   doc.add_paragraph(f"Fan: {fan}")
    if sinf:  doc.add_paragraph(f"Sinf: {sinf}")
    doc.add_paragraph()

    headers = ["O'quvchi", "Sinf", "Fan", "Jami", "To'g'ri", "Natija"]
    data    = [(r[0], r[1], r[2], r[3], r[4], f"{r[5]}%") for r in rows]
    add_table(doc, headers, data)

    doc.add_paragraph()
    if rows:
        avg = sum(r[5] or 0 for r in rows) / len(rows)
        p   = doc.add_paragraph()
        run = p.add_run(f"O'rtacha natija: {avg:.1f}%")
        run.bold = True; run.font.size = Pt(12)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# ══════════════════════════════════════
# 2. KITOB BO'LIMLARI WORD (10-20 betlik)
# ══════════════════════════════════════
def book_to_docx_chunks(book_id: int, pages_per_chunk: int = 15) -> list:
    """
    Kitobni Word hujjatlariga bo'ladi.
    Har chunk pages_per_chunk sahifalik.
    Qaytadi: [(filename, BytesIO), ...]
    """
    conn = db(); cur = conn.cursor()
    cur.execute("SELECT title, fan, sinf, muallif FROM books WHERE id=%s", (book_id,))
    book = cur.fetchone()
    if not book: return []
    title, fan, sinf, muallif = book

    cur.execute("""
        SELECT bs.title, bc.matn, bc.latex, bc.chunk_type, bc.page_num,
               bi.file_id
        FROM book_chunks bc
        JOIN book_sections bs ON bs.id = bc.section_id
        LEFT JOIN book_images bi ON bi.section_id = bc.section_id
          AND bi.page_num = bc.page_num
        WHERE bc.book_id = %s
        ORDER BY bc.page_num, bc.id
    """, (book_id,))
    chunks = cur.fetchall()
    cur.close(); conn.close()

    if not chunks: return []

    # Sahifalarga bo'lamiz
    pages   = {}
    for ch in chunks:
        pg = ch[4] or 1
        pages.setdefault(pg, []).append(ch)

    page_nums   = sorted(pages.keys())
    chunk_files = []
    chunk_idx   = 1

    for i in range(0, len(page_nums), pages_per_chunk):
        chunk_pages = page_nums[i:i+pages_per_chunk]
        doc = Document()
        doc.core_properties.author = "SamTM"

        # Sarlavha
        p_from = chunk_pages[0]
        p_to   = chunk_pages[-1]
        set_heading(doc, title, 1, (30,80,120))
        doc.add_paragraph(f"{fan} | {sinf}-sinf | {muallif}")
        doc.add_paragraph(f"Qism {chunk_idx}: {p_from}-{p_to} betlar")
        doc.add_page_break()

        cur_section = None
        for pg in chunk_pages:
            for (sec_title, matn, latex, ctype, page_num, img_fid) in pages.get(pg, []):
                # Yangi bo'lim sarlavhasi
                if sec_title != cur_section:
                    set_heading(doc, sec_title, 2, (46,134,171))
                    cur_section = sec_title

                if ctype == "formula" and latex:
                    p = doc.add_paragraph()
                    r = p.add_run(latex)
                    r.font.name = "Courier New"
                    r.font.size = Pt(11)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif matn:
                    doc.add_paragraph(matn)

        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        fname = f"{title[:20]}_qism{chunk_idx}_{p_from}-{p_to}.docx"
        chunk_files.append((fname, buf))
        chunk_idx += 1

    return chunk_files

# ══════════════════════════════════════
# 3. DARS REJASI WORD
# ══════════════════════════════════════
def lesson_plan_docx(sinf: str, fan: str, weeks: int = 4) -> io.BytesIO:
    conn = db(); cur = conn.cursor()
    cur.execute("""
        SELECT d.mavzu_name, d.kichik_name, d.topic_code,
               EXISTS(SELECT 1 FROM teacher_lessons tl WHERE tl.topic_code=d.topic_code),
               EXISTS(SELECT 1 FROM generated_tests gt WHERE gt.topic_code=d.topic_code)
        FROM dts_tree d
        WHERE d.grade=%s AND d.subject_name=%s AND d.is_deleted=FALSE
        ORDER BY d.topic_code LIMIT %s
    """, (sinf, fan, weeks*5))
    topics = cur.fetchall()
    cur.close(); conn.close()

    doc = Document()
    set_heading(doc, f"{sinf}-sinf | {fan}", 1, (30,80,120))
    set_heading(doc, f"{weeks} haftalik dars rejasi", 2)
    doc.add_paragraph(f"Tuzildi: {datetime.now().strftime('%d.%m.%Y')}")
    doc.add_paragraph()

    days_uz = ["Dushanba","Seshanba","Chorshanba","Payshanba","Juma"]
    idx = 0
    for w in range(1, weeks+1):
        set_heading(doc, f"{w}-hafta", 3, (46,134,171))
        rows = []
        for d_name in days_uz:
            if idx >= len(topics): break
            t = topics[idx]
            rows.append([d_name, t[0], t[1],
                         "✅" if t[3] else "❌",
                         "✅" if t[4] else "❌"])
            idx += 1
        add_table(doc, ["Kun","Mavzu","Kichik mavzu","Dars","Test"], rows)
        doc.add_paragraph()

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# ══════════════════════════════════════
# 4. ZIP ARXIV YARATISH
# ══════════════════════════════════════
def create_zip(files: list, zip_name: str = "arxiv.zip") -> io.BytesIO:
    """
    files = [(filename, BytesIO), ...]
    Qaytadi: ZIP fayli BytesIO
    """
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for fname, fbuf in files:
            fbuf.seek(0)
            zf.writestr(fname, fbuf.read())
    buf.seek(0)
    return buf

def create_book_archive(book_id: int, pages_per_chunk: int = 15) -> tuple:
    """
    Kitobni bo'laklarga bo'lib ZIP arxiv yaratadi.
    Qaytadi: (zip_buf, info_text)
    """
    chunks = book_to_docx_chunks(book_id, pages_per_chunk)
    if not chunks:
        return None, "❌ Kitob topilmadi yoki bo'sh."

    conn = db(); cur = conn.cursor()
    cur.execute("SELECT title, fan, sinf FROM books WHERE id=%s", (book_id,))
    book = cur.fetchone(); cur.close(); conn.close()
    title = book[0] if book else f"Kitob_{book_id}"

    zip_buf  = create_zip(chunks, f"{title}.zip")
    info     = (
        f"📦 {title}\n"
        f"📚 {len(chunks)} ta Word fayl\n"
        f"📄 Har biri ~{pages_per_chunk} bet\n"
        f"🗜 ZIP arxivda"
    )
    return zip_buf, info

# ══════════════════════════════════════
# 5. TO'LIQ PAKET (barcha hujjatlar)
# ══════════════════════════════════════
def full_package(sinf: str, fan: str, days: int = 30) -> io.BytesIO:
    """
    Bir tugma bilan: test natijalari + dars rejasi + taraqqiyot
    Hammasi bitta ZIP da
    """
    from jadval_generator import (test_results_excel, lesson_plan_excel,
                                   student_progress_excel)
    files = []
    try:
        files.append(("test_natijalari.xlsx",    test_results_excel(fan, sinf, days)))
        files.append(("dars_rejasi.xlsx",         lesson_plan_excel(sinf, fan)))
        files.append(("student_taraqqiyot.xlsx",  student_progress_excel(sinf)))
        files.append(("test_natijalari.docx",     test_results_docx(fan, sinf, days)))
        files.append(("dars_rejasi.docx",          lesson_plan_docx(sinf, fan)))
    except Exception as e:
        print(f"Paket xato: {e}")

    return create_zip(files, f"paket_{sinf}sinf_{fan[:10]}.zip")
