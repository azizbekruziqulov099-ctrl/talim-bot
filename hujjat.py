"""hujjat.py — Rasmdan matn o'qib (OCR), chiroyli PDF yoki Word hujjat yaratadi.

OCR — Gemini Vision API orqali (GEMINI_API_KEY talab qilinadi).
Chiqish — PDF (fpdf2) yoki DOCX (python-docx).

O'zbek harflari (o', g' kabi) uchun Unicode shrift kerak — buni bir nechta
manbadan (matplotlib bundle, tizim yo'llari) qidirib topamiz, birortasi
topilmasa oddiy shriftga tushamiz (maxsus harflar noto'g'ri chiqishi mumkin,
lekin dastur qulamaydi).
"""
import os
import json
import base64
import urllib.request
import urllib.error


# ═══════════════ 1. RASMDAN MATN O'QISH (OCR) ═══════════════

def matn_ol_rasmdan(rasm_yol):
    """Gemini Vision orqali rasmdagi matnni o'qiydi, tuzilmasini
    (paragraflar) saqlab qaytaradi. (matn, xato)"""
    api_key = os.getenv("GEMINI_API_KEY", "")
    if not api_key:
        return (None, "GEMINI_API_KEY sozlanmagan (Railway muhitida yo'q)")

    try:
        with open(rasm_yol, "rb") as f:
            rasm_bytes = f.read()
        rasm_b64 = base64.b64encode(rasm_bytes).decode("utf-8")

        ext = rasm_yol.lower().rsplit(".", 1)[-1]
        mime = {"png": "image/png", "webp": "image/webp"}.get(ext, "image/jpeg")

        prompt = (
            "Bu rasmdagi barcha matnni AYNAN, so'zma-so'z o'qib yoz. "
            "Paragraflar orasida bo'sh qator qoldir. Sarlavha bo'lsa "
            "birinchi qatorda alohida yoz. Rasmda yo'q narsa qo'shma, "
            "izoh-tushuntirish yozma, FAQAT rasmdagi matnni qaytar."
        )

        url = ("https://generativelanguage.googleapis.com/v1beta/"
               f"models/gemini-2.5-flash:generateContent?key={api_key}")
        body = json.dumps({
            "contents": [{
                "parts": [
                    {"text": prompt},
                    {"inline_data": {"mime_type": mime, "data": rasm_b64}}
                ]
            }]
        }).encode("utf-8")
        req = urllib.request.Request(url, data=body,
                                     headers={"Content-Type": "application/json"})
        with urllib.request.urlopen(req, timeout=60) as resp:
            natija = json.loads(resp.read().decode("utf-8"))
        matn = natija["candidates"][0]["content"]["parts"][0]["text"].strip()
        if not matn:
            return (None, "Rasmda matn topilmadi")
        return (matn, None)
    except urllib.error.HTTPError as e:
        return (None, f"Gemini xato ({e.code}): {e.read()[:200]}")
    except Exception as e:
        return (None, str(e)[:300])


# ═══════════════ 2. UNICODE SHRIFT QIDIRISH ═══════════════

_SHRIFT_KESH = None

def _unicode_shrift_topish():
    """O'zbek harflarini (o', g') to'g'ri ko'rsatadigan shrift qidiradi.
    Bir necha manbadan — birortasi ishlasa bo'ldi. Natija keshlanadi."""
    global _SHRIFT_KESH
    if _SHRIFT_KESH is not None:
        return _SHRIFT_KESH or None

    nomzodlar = []
    try:
        import matplotlib
        nomzodlar.append(os.path.join(matplotlib.get_data_path(), "fonts", "ttf", "DejaVuSans.ttf"))
    except Exception:
        pass
    nomzodlar += [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/dejavu/DejaVuSans.ttf",
        "/usr/local/share/fonts/DejaVuSans.ttf",
    ]
    for yol in nomzodlar:
        if os.path.exists(yol):
            _SHRIFT_KESH = yol
            return yol
    _SHRIFT_KESH = ""
    return None


# ═══════════════ 3. HUJJAT YARATISH ═══════════════

def matndan_pdf(matn, chiqish_yol, sarlavha=None):
    """Matndan chiroyli formatlangan PDF yaratadi. (muvaffaqiyat, xato)"""
    from fpdf import FPDF
    try:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        shrift_yol = _unicode_shrift_topish()
        if shrift_yol:
            pdf.add_font("Unicode", "", shrift_yol)
            asosiy_shrift = "Unicode"
        else:
            asosiy_shrift = "Helvetica"

        # MUHIM: multi_cell() dan keyin kursor CHAP CHEKKAGA qaytmaydi —
        # o'ng chekkada qolib ketadi. Har chaqiruvdan oldin qo'lda
        # qaytarish shart, aks holda "joy yo'q" xatosi chiqadi.
        if sarlavha:
            pdf.set_font(asosiy_shrift, size=16)
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(0, 10, sarlavha)
            pdf.set_x(pdf.l_margin)
            pdf.ln(4)

        pdf.set_font(asosiy_shrift, size=12)
        for band in matn.split("\n"):
            pdf.set_x(pdf.l_margin)
            if band.strip():
                pdf.multi_cell(0, 8, band)
            else:
                pdf.ln(4)

        pdf.output(chiqish_yol)
        if not os.path.exists(chiqish_yol):
            return (False, "PDF fayli yaratilmadi")
        return (True, None)
    except Exception as e:
        return (False, str(e)[:300])


# ═══════════════ 4. HAQIQIY SKANER (chegara aniqlash + tekislash + oq-qora) ═══════════════
# CamScanner/Adobe Scan uslubida: hujjat chegarasini AVTOMATIK topadi,
# perspektivani tekislaydi, skaner ko'rinishiga (oq fon, aniq matn) keltiradi.
# ESLATMA: Telegram orqali "4 burchakni qo'lda surish" TEXNIK JIHATDAN
# IMKONSIZ (bot tugmalari bilan bunday interfeys qurib bo'lmaydi) — shuning
# uchun avtomatik aniqlashga tayanamiz, u yetarlicha ishonchli ishlaydi.

def _nuqtalarni_tartibla(nuqtalar):
    """4 nuqtani [yuqori-chap, yuqori-o'ng, past-o'ng, past-chap] tartibiga soladi."""
    import numpy as np
    natija = np.zeros((4, 2), dtype="float32")
    yigindi = nuqtalar.sum(axis=1)
    natija[0] = nuqtalar[np.argmin(yigindi)]
    natija[2] = nuqtalar[np.argmax(yigindi)]
    farq = np.diff(nuqtalar, axis=1)
    natija[1] = nuqtalar[np.argmin(farq)]
    natija[3] = nuqtalar[np.argmax(farq)]
    return natija


def rasmni_skanla(rasm_yol, chiqish_yol, oq_qora=True, kesish_kengaytma=0.0):
    """Hujjat chegarasini avtomatik topib kesadi, perspektivani tekislaydi,
    skaner ko'rinishiga (oq fon) keltiradi. HECH QACHON butunlay rad etmaydi —
    aniq 4-burchak topilmasa, minAreaRect (eng kichik o'rovchi to'rtburchak)
    bilan taxminiy chegara oladi; u ham bo'lmasa, butun rasmni ishlatadi.

    kesish_kengaytma: -0.15..0.15 oralig'i — manfiy=torroq kesadi (chetlardan
        ko'proq olib tashlaydi), musbat=kengroq kesadi (chetlarni ko'proq
        qoldiradi). "Sozlamadan to'g'irlash" uchun. (muvaffaqiyat, xato)"""
    try:
        import cv2
        import numpy as np

        img = cv2.imread(rasm_yol)
        if img is None:
            return (False, "Rasm o'qilmadi")
        orig = img.copy()
        h, w = img.shape[:2]

        olchov = 800 / max(h, w) if max(h, w) > 800 else 1.0
        kichik = cv2.resize(img, None, fx=olchov, fy=olchov)
        kulrang = cv2.cvtColor(kichik, cv2.COLOR_BGR2GRAY)
        xira = cv2.GaussianBlur(kulrang, (5, 5), 0)
        qirralar = cv2.Canny(xira, 50, 150)
        qirralar = cv2.dilate(qirralar, None, iterations=2)

        konturlar, _ = cv2.findContours(qirralar, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        konturlar = sorted(konturlar, key=cv2.contourArea, reverse=True)[:5]
        rasm_yuzasi = kichik.shape[0] * kichik.shape[1]

        nuqtalar = None

        # 1-urinish: aniq 4-burchakli kontur (eng ishonchli, yaqindan olingan rasm)
        for k in konturlar:
            if cv2.contourArea(k) < rasm_yuzasi * 0.15:
                continue
            perimetr = cv2.arcLength(k, True)
            taxmin = cv2.approxPolyDP(k, 0.02 * perimetr, True)
            if len(taxmin) == 4:
                nuqtalar = (taxmin.reshape(4, 2) / olchov).astype("float32")
                break

        # 2-urinish: aniq 4-burchak topilmasa — minAreaRect (uzoqdan olingan,
        # kichikroq hujjat, yoki qirralar to'liq tekis chiqmagan holatlar uchun)
        if nuqtalar is None and konturlar:
            eng_katta = konturlar[0]
            if cv2.contourArea(eng_katta) > rasm_yuzasi * 0.02:
                rect = cv2.minAreaRect(eng_katta)
                qutilar = cv2.boxPoints(rect)
                nuqtalar = (qutilar / olchov).astype("float32")

        if nuqtalar is None:
            # Hech narsa topilmadi — butun rasmni ishlatamiz (rad ETMAYMIZ)
            tekislangan = orig
        else:
            nuqtalar = _nuqtalarni_tartibla(nuqtalar)

            # Kesish kengaytmasi — "sozlamadan to'g'irlash" uchun
            if kesish_kengaytma:
                markaz = nuqtalar.mean(axis=0)
                nuqtalar = markaz + (nuqtalar - markaz) * (1.0 + kesish_kengaytma)

            (tl, tr, br, bl) = nuqtalar
            kenglik = max(int(np.linalg.norm(br - bl)), int(np.linalg.norm(tr - tl)))
            balandlik = max(int(np.linalg.norm(tr - br)), int(np.linalg.norm(tl - bl)))
            if kenglik < 10 or balandlik < 10:
                tekislangan = orig
            else:
                maqsad = np.array([[0, 0], [kenglik - 1, 0],
                                   [kenglik - 1, balandlik - 1], [0, balandlik - 1]],
                                  dtype="float32")
                matritsa = cv2.getPerspectiveTransform(nuqtalar, maqsad)
                tekislangan = cv2.warpPerspective(orig, matritsa, (kenglik, balandlik))

        if oq_qora:
            kulrang2 = cv2.cvtColor(tekislangan, cv2.COLOR_BGR2GRAY)
            natija_rasm = cv2.adaptiveThreshold(
                kulrang2, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                cv2.THRESH_BINARY, 15, 10)
        else:
            # Rangli, lekin yorqinlik/kontrastni yaxshilaymiz
            lab = cv2.cvtColor(tekislangan, cv2.COLOR_BGR2LAB)
            l, a, b = cv2.split(lab)
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
            l = clahe.apply(l)
            natija_rasm = cv2.cvtColor(cv2.merge((l, a, b)), cv2.COLOR_LAB2BGR)

        cv2.imwrite(chiqish_yol, natija_rasm)
        if not os.path.exists(chiqish_yol) or os.path.getsize(chiqish_yol) == 0:
            return (False, "Skanlangan rasm saqlanmadi")
        return (True, None)
    except Exception as e:
        return (False, str(e)[:300])


def rasmlardan_skan_pdf(rasm_yollari, chiqish_yol, oq_qora=True, kesish_kengaytma=0.0):
    """Bir yoki bir nechta rasmni SKANLAB (chegara+tekislash+oq-qora),
    ko'p sahifali PDF yaratadi. (muvaffaqiyat, xato)"""
    from PIL import Image
    try:
        if not rasm_yollari:
            return (False, "Rasm berilmagan")

        skanlangan_yollar = []
        papka = os.path.dirname(chiqish_yol) or "."
        for i, yol in enumerate(rasm_yollari):
            skan_yol = os.path.join(papka, f"_skan_{i}.jpg")
            ok, xato = rasmni_skanla(yol, skan_yol, oq_qora=oq_qora,
                                     kesish_kengaytma=kesish_kengaytma)
            if not ok:
                # Bitta rasm muvaffaqiyatsiz bo'lsa — asl rasmni ishlatamiz, to'xtamaymiz
                skan_yol = yol
            skanlangan_yollar.append(skan_yol)

        rasmlar = [Image.open(y).convert("RGB") for y in skanlangan_yollar]
        rasmlar[0].save(chiqish_yol, save_all=True, append_images=rasmlar[1:])
        if not os.path.exists(chiqish_yol) or os.path.getsize(chiqish_yol) == 0:
            return (False, "PDF yaratilmadi")
        return (True, None)
    except Exception as e:
        return (False, str(e)[:300])

def matndan_docx(matn, chiqish_yol, sarlavha=None):
    """Matndan Word hujjat yaratadi. (muvaffaqiyat, xato)"""
    from docx import Document
    try:
        doc = Document()
        if sarlavha:
            doc.add_heading(sarlavha, level=1)
        for band in matn.split("\n"):
            if band.strip():
                doc.add_paragraph(band)
        doc.save(chiqish_yol)
        if not os.path.exists(chiqish_yol):
            return (False, "Word fayli yaratilmadi")
        return (True, None)
    except Exception as e:
        return (False, str(e)[:300])
